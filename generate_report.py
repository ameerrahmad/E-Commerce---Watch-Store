from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ChronoLux — React TypeScript Conversion")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x0B, 0x0F, 0x1A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Converting HTML/CSS to React Environment")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Class Task — React.js with TypeScript")
r.font.size = Pt(14)

doc.add_paragraph("")
doc.add_paragraph("")

info_lines = [
    ("Submitted By:", "Ameer Ahmad"),
    ("Roll Number:", "24011519-143"),
    ("Email:", "ameerr.ahmad2005@gmail.com"),
    ("Repository:", "github.com/ameerrahmad/E-Commerce---Watch-Store"),
    ("Date:", "May 2026"),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.size = Pt(12)
    r = p.add_run(value)
    r.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Task Description",
    "2. Tools & Technologies",
    "3. Step-by-Step Process",
    "4. Project Structure",
    "5. HTML to JSX Conversion — Key Changes",
    "6. Code Comparison (Before & After)",
    "7. Output / Screenshots",
    "8. Conclusion",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. TASK DESCRIPTION
# ============================================================
doc.add_heading("1. Task Description", level=1)
doc.add_paragraph(
    "The objective of this class task is to convert the existing ChronoLux luxury watch store "
    "website (previously built with plain HTML, CSS, and JavaScript in Assignment #1) into a "
    "React application using TypeScript. This conversion involves:"
)
tasks = [
    "Creating a new React application using Vite with TypeScript template",
    "Keeping the default Vite folder structure unchanged",
    "Copying the CSS from globalStyles.css into the React app's index.css",
    "Copying the HTML from index.html into App.tsx and converting it to valid JSX",
    "Fixing all JSX syntax errors that arise from the HTML-to-JSX conversion",
    "No Tailwind CSS — only vanilla CSS from Assignment #1",
]
for t in tasks:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph(
    "This task is a direct conversion exercise — the goal is to get the same visual output "
    "running inside a React environment without changing the design or adding new features."
)

doc.add_page_break()

# ============================================================
# 2. TOOLS & TECHNOLOGIES
# ============================================================
doc.add_heading("2. Tools & Technologies", level=1)
techs = [
    ("React 19", "JavaScript library for building user interfaces"),
    ("TypeScript", "Typed superset of JavaScript for type safety"),
    ("Vite", "Fast build tool and development server"),
    ("CSS3", "Vanilla CSS with custom properties (no Tailwind)"),
    ("Node.js & npm", "Runtime and package manager"),
    ("VS Code", "Code editor"),
    ("Google Fonts", "Cormorant Garamond and Inter typefaces"),
]
for name, desc in techs:
    p = doc.add_paragraph()
    r = p.add_run(name + ": ")
    r.bold = True
    p.add_run(desc)

doc.add_page_break()

# ============================================================
# 3. STEP-BY-STEP PROCESS
# ============================================================
doc.add_heading("3. Step-by-Step Process", level=1)

doc.add_heading("Step 1: Create the React App", level=2)
doc.add_paragraph("The React application was created using Vite with the react-ts template:")
p = doc.add_paragraph()
r = p.add_run("npm create vite@latest ./ -- --template react-ts")
r.font.name = "Consolas"
r.font.size = Pt(10)
doc.add_paragraph(
    "This scaffolds a React + TypeScript project with all necessary configuration files "
    "(tsconfig.json, vite.config.ts, eslint.config.js) and the src/ folder structure."
)

doc.add_heading("Step 2: Install Dependencies", level=2)
p = doc.add_paragraph()
r = p.add_run("npm install")
r.font.name = "Consolas"
r.font.size = Pt(10)
doc.add_paragraph("This installs React, ReactDOM, TypeScript, and Vite dependencies.")

doc.add_heading("Step 3: Copy CSS into index.css", level=2)
doc.add_paragraph(
    "The entire content of globalStyles.css (1429 lines) from Assignment #1 was copied into "
    "src/index.css, replacing the default Vite styles. This file is already imported in main.tsx "
    "so all styles are automatically applied."
)

doc.add_heading("Step 4: Copy Images to public/", level=2)
doc.add_paragraph(
    "The sources/ folder containing all watch images was copied into the public/ directory "
    "so that images are accessible via paths like 'sources/image3.png' in the React app."
)

doc.add_heading("Step 5: Add Google Fonts to index.html", level=2)
doc.add_paragraph(
    "The Vite index.html (in the project root, not src/) was updated to include Google Fonts "
    "preconnect links and the stylesheet for Cormorant Garamond and Inter fonts."
)

doc.add_heading("Step 6: Convert HTML to JSX in App.tsx", level=2)
doc.add_paragraph(
    "The HTML from index.html was pasted into App.tsx and all syntax errors were fixed by "
    "applying the JSX conversion rules (detailed in Section 5). The JavaScript functions "
    "(cart management, dark mode, toast notifications) were moved inside the React component."
)

doc.add_heading("Step 7: Clear App.css", level=2)
doc.add_paragraph(
    "The default App.css was emptied since all styles come from index.css (globalStyles.css)."
)

doc.add_heading("Step 8: Verify", level=2)
doc.add_paragraph("TypeScript compilation check was run with zero errors:")
p = doc.add_paragraph()
r = p.add_run("npx tsc --noEmit    # 0 errors")
r.font.name = "Consolas"
r.font.size = Pt(10)
doc.add_paragraph("")
doc.add_paragraph("Development server was started successfully:")
p = doc.add_paragraph()
r = p.add_run("npm run dev         # Running at http://localhost:5173/")
r.font.name = "Consolas"
r.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 4. PROJECT STRUCTURE
# ============================================================
doc.add_heading("4. Project Structure", level=1)
doc.add_paragraph("The default Vite folder structure was kept unchanged:")

structure = """chronolux-react/
|-- index.html              (Vite entry + Google Fonts)
|-- package.json
|-- tsconfig.json
|-- tsconfig.app.json
|-- tsconfig.node.json
|-- vite.config.ts
|-- eslint.config.js
|-- public/
|   |-- favicon.svg
|   |-- sources/            (watch images copied from Assignment#1)
|       |-- image1.png
|       |-- image2.png
|       |-- image3.png
|       |-- image5.png
|       |-- image6.png
|-- src/
|   |-- main.tsx            (React entry point - unchanged)
|   |-- App.tsx             (Converted from index.html)
|   |-- App.css             (Emptied - styles in index.css)
|   |-- index.css           (Copied from globalStyles.css)
|   |-- assets/
|-- node_modules/"""

p = doc.add_paragraph()
r = p.add_run(structure)
r.font.size = Pt(9)
r.font.name = "Consolas"

doc.add_page_break()

# ============================================================
# 5. HTML TO JSX CONVERSION — KEY CHANGES
# ============================================================
doc.add_heading("5. HTML to JSX Conversion — Key Changes", level=1)
doc.add_paragraph(
    "When converting HTML to JSX (React), several syntax changes are required because "
    "JSX is JavaScript, not HTML. Below are all the changes that were applied:"
)

changes = [
    ("class → className",
     'HTML uses class=\"hero\" but JSX uses className=\"hero\" because class is a reserved word in JavaScript.'),
    ("for → htmlFor",
     'HTML uses for=\"email\" on labels but JSX uses htmlFor=\"email\" because for is a reserved word in JavaScript.'),
    ("onclick → onClick",
     'HTML event handlers are lowercase (onclick) but JSX uses camelCase (onClick). Also, inline string handlers become arrow functions: onClick={() => addToCart(...)}'),
    ("onsubmit → onSubmit",
     'Same camelCase rule. The handler becomes: onSubmit={(e) => { e.preventDefault(); showToast(...); }}'),
    ("style strings → style objects",
     'HTML: style=\"background-image:url(...)\" becomes JSX: style={{ backgroundImage: \"url(...)\" }}. CSS properties use camelCase and values are strings.'),
    ("HTML comments → JSX comments",
     '<!-- comment --> becomes {/* comment */} because JSX uses JavaScript comment syntax inside curly braces.'),
    ("Self-closing tags",
     'Tags like <img>, <input>, <br> must be self-closed in JSX: <img />, <input />, <br />'),
    ("Fragment wrapper",
     'React components must return a single root element. The entire page content is wrapped in <> ... </> (React Fragment).'),
    ("No <script> tags",
     'JavaScript code is moved inside the component function instead of using <script> tags.'),
    ("&amp; entity",
     'HTML entities like &amp; can be used directly or as {\"&\"} in JSX.'),
]

for title, desc in changes:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    doc.add_paragraph(desc)
    doc.add_paragraph("")

doc.add_page_break()

# ============================================================
# 6. CODE COMPARISON
# ============================================================
doc.add_heading("6. Code Comparison (Before & After)", level=1)

doc.add_heading("Example 1: Hero Section", level=2)
doc.add_paragraph("BEFORE (HTML — index.html):")
before1 = """<section class="hero">
  <div class="hero-bg"
       style="background-image:url('sources/image1.png');">
  </div>
  <div class="hero-content">
    <h1>Time Is Your
        <span class="accent-text">Masterpiece</span>
    </h1>
  </div>
</section>"""
p = doc.add_paragraph()
r = p.add_run(before1)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_paragraph("")
doc.add_paragraph("AFTER (JSX — App.tsx):")
after1 = """<section className="hero">
  <div className="hero-bg"
       style={{ backgroundImage: "url('sources/image1.png')" }}>
  </div>
  <div className="hero-content">
    <h1>Time Is Your
        <span className="accent-text">Masterpiece</span>
    </h1>
  </div>
</section>"""
p = doc.add_paragraph()
r = p.add_run(after1)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_heading("Example 2: Button with Event Handler", level=2)
doc.add_paragraph("BEFORE (HTML):")
before2 = """<button class="sm-button border-button"
  onclick="addToCart('prod-1','Chronograph Pro',
           2850,'sources/image3.png')">
  Add to Cart
</button>"""
p = doc.add_paragraph()
r = p.add_run(before2)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_paragraph("")
doc.add_paragraph("AFTER (JSX):")
after2 = """<button className="sm-button border-button"
  onClick={() => addToCart('prod-1',
    'Chronograph Pro', 2850, 'sources/image3.png')}>
  Add to Cart
</button>"""
p = doc.add_paragraph()
r = p.add_run(after2)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_heading("Example 3: Form with onSubmit", level=2)
doc.add_paragraph("BEFORE (HTML):")
before3 = """<form onsubmit="event.preventDefault();
  showToast('Welcome!');" class="newsletter-form">"""
p = doc.add_paragraph()
r = p.add_run(before3)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_paragraph("")
doc.add_paragraph("AFTER (JSX):")
after3 = """<form onSubmit={(e) => {
    e.preventDefault();
    showToast('Welcome!');
  }} className="newsletter-form">"""
p = doc.add_paragraph()
r = p.add_run(after3)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_heading("Example 4: JavaScript Functions", level=2)
doc.add_paragraph("BEFORE (HTML — separate <script> tag):")
before4 = """<script src="jsFile/jsFunctions.js"></script>"""
p = doc.add_paragraph()
r = p.add_run(before4)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_paragraph("")
doc.add_paragraph("AFTER (Inside React component):")
after4 = """function App() {
  const showToast = (message: string) => {
    const toast = document.createElement('div');
    toast.className = 'toast';
    toast.textContent = message;
    document.body.appendChild(toast);
    setTimeout(() => toast.remove(), 3000);
  };

  const addToCart = (id: string, name: string,
    price: number, image: string) => {
    const cart = getCart();
    // ... cart logic
    showToast(name + ' added to cart!');
  };

  return ( /* JSX here */ );
}
export default App"""
p = doc.add_paragraph()
r = p.add_run(after4)
r.font.size = Pt(8.5)
r.font.name = "Consolas"

doc.add_page_break()

# ============================================================
# 7. OUTPUT / SCREENSHOTS
# ============================================================
doc.add_heading("7. Output / Screenshots", level=1)
doc.add_paragraph("The React application runs successfully at http://localhost:5173/ with the same visual output as the original HTML version.")

import os
screenshots = [
    ("C:\\Users\\Shining star\\.gemini\\antigravity\\brain\\7067c55b-83a2-45a3-88ed-eb023d62b7ac\\home_hero_section_1778267302986.png",
     "React App — Homepage Hero Section (identical to HTML version)"),
    ("C:\\Users\\Shining star\\.gemini\\antigravity\\brain\\7067c55b-83a2-45a3-88ed-eb023d62b7ac\\home_card_section_1778267321174.png",
     "React App — Featured Timepieces Cards"),
]

for img_path, caption in screenshots:
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        doc.add_paragraph("")

doc.add_paragraph("")
doc.add_paragraph("TypeScript Compilation Result:")
p = doc.add_paragraph()
r = p.add_run("$ npx tsc --noEmit\n(no errors — exit code 0)")
r.font.name = "Consolas"
r.font.size = Pt(10)

doc.add_paragraph("")
doc.add_paragraph("Vite Dev Server Output:")
p = doc.add_paragraph()
r = p.add_run("VITE v8.0.11  ready in 1252 ms\n  > Local:   http://localhost:5173/")
r.font.name = "Consolas"
r.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 8. CONCLUSION
# ============================================================
doc.add_heading("8. Conclusion", level=1)
doc.add_paragraph(
    "The ChronoLux website from Assignment #1 was successfully converted from a plain "
    "HTML/CSS/JavaScript project into a React application using TypeScript. The key outcomes are:"
)
conclusions = [
    "React app created using Vite with react-ts template",
    "Default Vite folder structure preserved — no structural changes",
    "All 1429 lines of CSS from globalStyles.css transferred to index.css",
    "index.html content converted to valid JSX in App.tsx with zero TypeScript errors",
    "All HTML-to-JSX syntax issues resolved (className, onClick, style objects, etc.)",
    "JavaScript functions moved inside the React component with proper TypeScript types",
    "No Tailwind CSS used — only vanilla CSS from Assignment #1",
    "The visual output is identical to the original HTML version",
]
for c in conclusions:
    doc.add_paragraph(c, style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph(
    "This exercise demonstrates how an existing static website can be migrated into a modern "
    "React environment while maintaining the same design and functionality. The React version "
    "provides a foundation for future enhancements like component splitting, routing with "
    "React Router, and state management."
)

# SAVE
output_path = r"c:\Users\Shining star\Documents\E-Commerce---Watch-Store\chronolux-react\ChronoLux_React_Task_Report.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
